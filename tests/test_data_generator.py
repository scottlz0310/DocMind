"""
テストデータジェネレーター

様々なドキュメントタイプのテストデータを生成するユーティリティ
"""

import json
import random
import shutil
import tempfile
from datetime import datetime
from pathlib import Path

# PDFテスト用
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Wordテスト用
try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excelテスト用
try:
    from openpyxl import Workbook

    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


class TestDataGenerator:
    """テストデータ生成クラス"""

    def __init__(self, base_dir: str = None):
        """
        初期化

        Args:
            base_dir: テストデータを生成するベースディレクトリ
        """
        self.base_dir = Path(base_dir) if base_dir else Path(tempfile.mkdtemp())
        self.base_dir.mkdir(exist_ok=True)

        # サンプルテキストデータ
        self.sample_texts = [
            "これは日本語のサンプルテキストです。検索機能のテストに使用されます。",
            "Python programming language is widely used for data science and machine learning.",
            "人工知能と機械学習の技術が急速に発展しています。",
            "Document search and retrieval systems are essential for knowledge management.",
            "データベース設計とパフォーマンス最適化について説明します。",
            "Natural language processing enables computers to understand human language.",
            "ソフトウェア開発におけるテスト駆動開発の重要性について。",
            "Cloud computing and distributed systems architecture patterns.",
        ]

        # ファイル名のサンプル
        self.sample_filenames = [
            "技術仕様書",
            "設計書",
            "要件定義",
            "テスト計画",
            "ユーザーマニュアル",
            "API仕様",
            "データベース設計",
            "システム構成",
            "運用手順",
            "障害対応",
        ]

    def generate_text_content(self, length: int = 500) -> str:
        """
        指定された長さのテキストコンテンツを生成

        Args:
            length: 生成するテキストの文字数

        Returns:
            生成されたテキスト
        """
        content = ""
        while len(content) < length:
            content += random.choice(self.sample_texts) + "\n\n"

        return content[:length]

    def create_text_files(self, count: int = 10) -> list[Path]:
        """
        テキストファイルを生成

        Args:
            count: 生成するファイル数

        Returns:
            生成されたファイルパスのリスト
        """
        text_dir = self.base_dir / "text_files"
        text_dir.mkdir(exist_ok=True)

        files = []
        for i in range(count):
            filename = f"{random.choice(self.sample_filenames)}_{i:03d}.txt"
            file_path = text_dir / filename

            content = self.generate_text_content(random.randint(100, 2000))
            file_path.write_text(content, encoding="utf-8")
            files.append(file_path)

        return files

    def create_markdown_files(self, count: int = 10) -> list[Path]:
        """
        Markdownファイルを生成

        Args:
            count: 生成するファイル数

        Returns:
            生成されたファイルパスのリスト
        """
        md_dir = self.base_dir / "markdown_files"
        md_dir.mkdir(exist_ok=True)

        files = []
        for i in range(count):
            filename = f"{random.choice(self.sample_filenames)}_{i:03d}.md"
            file_path = md_dir / filename

            # Markdownコンテンツを生成
            content = f"# {random.choice(self.sample_filenames)}\n\n"
            content += f"## 概要\n\n{self.generate_text_content(300)}\n\n"
            content += f"## 詳細\n\n{self.generate_text_content(500)}\n\n"
            content += "### サブセクション\n\n"
            content += f"- 項目1: {random.choice(self.sample_texts)}\n"
            content += f"- 項目2: {random.choice(self.sample_texts)}\n"
            content += f"- 項目3: {random.choice(self.sample_texts)}\n"

            file_path.write_text(content, encoding="utf-8")
            files.append(file_path)

        return files

    def create_pdf_files(self, count: int = 5) -> list[Path]:
        """
        PDFファイルを生成（reportlabが利用可能な場合）

        Args:
            count: 生成するファイル数

        Returns:
            生成されたファイルパスのリスト
        """
        if not PDF_AVAILABLE:
            return []

        pdf_dir = self.base_dir / "pdf_files"
        pdf_dir.mkdir(exist_ok=True)

        files = []
        for i in range(count):
            filename = f"{random.choice(self.sample_filenames)}_{i:03d}.pdf"
            file_path = pdf_dir / filename

            # PDFを作成
            c = canvas.Canvas(str(file_path), pagesize=letter)

            # 日本語フォントを設定（利用可能な場合）
            try:
                pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
                c.setFont("HeiseiKakuGo-W5", 12)
            except:
                c.setFont("Helvetica", 12)

            # コンテンツを追加
            y_position = 750
            title = f"{random.choice(self.sample_filenames)} {i+1}"
            c.drawString(100, y_position, title)

            y_position -= 50
            content_lines = self.generate_text_content(1000).split("\n")

            for line in content_lines[:30]:  # 最大30行
                if y_position < 100:
                    c.showPage()
                    y_position = 750
                    try:
                        c.setFont("HeiseiKakuGo-W5", 12)
                    except:
                        c.setFont("Helvetica", 12)

                # 日本語文字を含む場合の処理
                try:
                    c.drawString(100, y_position, line[:80])  # 80文字まで
                except:
                    # ASCII文字のみの場合
                    ascii_line = "".join(c for c in line if ord(c) < 128)
                    c.drawString(100, y_position, ascii_line[:80])

                y_position -= 20

            c.save()
            files.append(file_path)

        return files

    def create_word_files(self, count: int = 5) -> list[Path]:
        """
        Wordファイルを生成（python-docxが利用可能な場合）

        Args:
            count: 生成するファイル数

        Returns:
            生成されたファイルパスのリスト
        """
        if not DOCX_AVAILABLE:
            return []

        word_dir = self.base_dir / "word_files"
        word_dir.mkdir(exist_ok=True)

        files = []
        for i in range(count):
            filename = f"{random.choice(self.sample_filenames)}_{i:03d}.docx"
            file_path = word_dir / filename

            # Wordドキュメントを作成
            doc = Document()

            # タイトルを追加
            title = f"{random.choice(self.sample_filenames)} {i+1}"
            doc.add_heading(title, 0)

            # パラグラフを追加
            for _ in range(random.randint(3, 8)):
                paragraph_text = self.generate_text_content(random.randint(100, 300))
                doc.add_paragraph(paragraph_text)

            # サブヘッディングを追加
            doc.add_heading("詳細情報", level=1)
            for _ in range(random.randint(2, 5)):
                paragraph_text = self.generate_text_content(random.randint(50, 200))
                doc.add_paragraph(paragraph_text)

            doc.save(str(file_path))
            files.append(file_path)

        return files

    def create_excel_files(self, count: int = 5) -> list[Path]:
        """
        Excelファイルを生成（openpyxlが利用可能な場合）

        Args:
            count: 生成するファイル数

        Returns:
            生成されたファイルパスのリスト
        """
        if not EXCEL_AVAILABLE:
            return []

        excel_dir = self.base_dir / "excel_files"
        excel_dir.mkdir(exist_ok=True)

        files = []
        for i in range(count):
            filename = f"{random.choice(self.sample_filenames)}_{i:03d}.xlsx"
            file_path = excel_dir / filename

            # Excelワークブックを作成
            wb = Workbook()

            # デフォルトシートを削除
            wb.remove(wb.active)

            # 複数のシートを作成
            for sheet_num in range(random.randint(1, 3)):
                sheet_name = f"シート{sheet_num + 1}"
                ws = wb.create_sheet(title=sheet_name)

                # ヘッダー行を追加
                headers = ["項目", "説明", "値", "備考"]
                for col, header in enumerate(headers, 1):
                    ws.cell(row=1, column=col, value=header)

                # データ行を追加
                for row in range(2, random.randint(10, 50)):
                    ws.cell(row=row, column=1, value=f"項目{row-1}")
                    ws.cell(
                        row=row, column=2, value=random.choice(self.sample_texts)[:50]
                    )
                    ws.cell(row=row, column=3, value=random.randint(1, 1000))
                    ws.cell(
                        row=row, column=4, value=random.choice(self.sample_texts)[:30]
                    )

            wb.save(str(file_path))
            files.append(file_path)

        return files

    def create_comprehensive_test_dataset(
        self,
        text_count: int = 20,
        markdown_count: int = 15,
        pdf_count: int = 10,
        word_count: int = 10,
        excel_count: int = 8,
    ) -> dict[str, list[Path]]:
        """
        包括的なテストデータセットを作成

        Args:
            text_count: テキストファイル数
            markdown_count: Markdownファイル数
            pdf_count: PDFファイル数
            word_count: Wordファイル数
            excel_count: Excelファイル数

        Returns:
            ファイルタイプ別のファイルパス辞書
        """
        dataset = {}

        print("テストデータセットを生成中...")

        # 各タイプのファイルを生成
        dataset["text"] = self.create_text_files(text_count)
        print(f"テキストファイル {len(dataset['text'])} 個を生成")

        dataset["markdown"] = self.create_markdown_files(markdown_count)
        print(f"Markdownファイル {len(dataset['markdown'])} 個を生成")

        dataset["pdf"] = self.create_pdf_files(pdf_count)
        print(f"PDFファイル {len(dataset['pdf'])} 個を生成")

        dataset["word"] = self.create_word_files(word_count)
        print(f"Wordファイル {len(dataset['word'])} 個を生成")

        dataset["excel"] = self.create_excel_files(excel_count)
        print(f"Excelファイル {len(dataset['excel'])} 個を生成")

        # メタデータファイルを作成
        metadata = {
            "created_at": datetime.now().isoformat(),
            "total_files": sum(len(files) for files in dataset.values()),
            "file_counts": {k: len(v) for k, v in dataset.items()},
            "base_directory": str(self.base_dir),
        }

        metadata_path = self.base_dir / "dataset_metadata.json"
        metadata_path.write_text(
            json.dumps(metadata, indent=2, ensure_ascii=False), encoding="utf-8"
        )

        print(f"テストデータセット生成完了: 合計 {metadata['total_files']} ファイル")
        print(f"保存場所: {self.base_dir}")

        return dataset

    def cleanup(self):
        """生成されたテストデータを削除"""
        if self.base_dir.exists():
            shutil.rmtree(self.base_dir)
            print(f"テストデータを削除しました: {self.base_dir}")


def create_performance_test_dataset(size: str = "medium") -> TestDataGenerator:
    """
    パフォーマンステスト用のデータセットを作成

    Args:
        size: データセットサイズ ("small", "medium", "large")

    Returns:
        TestDataGeneratorインスタンス
    """
    size_configs = {
        "small": {
            "text_count": 50,
            "markdown_count": 30,
            "pdf_count": 10,
            "word_count": 10,
            "excel_count": 5,
        },
        "medium": {
            "text_count": 200,
            "markdown_count": 150,
            "pdf_count": 50,
            "word_count": 50,
            "excel_count": 25,
        },
        "large": {
            "text_count": 1000,
            "markdown_count": 750,
            "pdf_count": 200,
            "word_count": 200,
            "excel_count": 100,
        },
    }

    config = size_configs.get(size, size_configs["medium"])

    generator = TestDataGenerator()
    generator.create_comprehensive_test_dataset(**config)

    return generator


if __name__ == "__main__":
    # テスト実行例
    generator = create_performance_test_dataset("small")
    print("パフォーマンステスト用データセットを生成しました")
