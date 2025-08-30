"""
テストデータ生成クラス

包括的検証のための様々なテストデータを自動生成します。
"""

import csv
import json
import logging
import os
import random
import shutil
import tempfile
from dataclasses import dataclass
from datetime import datetime, timedelta
from typing import Any

# ドキュメント生成用ライブラリ（オプション）
try:
    import fitz  # PyMuPDF
    from docx import Document
    from openpyxl import Workbook

    HAS_DOCUMENT_LIBS = True
except ImportError:
    HAS_DOCUMENT_LIBS = False


@dataclass
class TestDatasetConfig:
    """テストデータセット設定"""

    dataset_name: str
    output_directory: str
    file_count: int = 100
    file_types: list[str] = None
    size_range_kb: tuple[int, int] = (1, 1000)
    content_language: str = "ja"
    include_corrupted: bool = False
    include_large_files: bool = False
    include_special_chars: bool = False


class TestDataGenerator:
    """
    テストデータ生成クラス

    包括的検証のための多様なテストデータを生成します。
    """

    def __init__(self):
        """テストデータ生成クラスの初期化"""
        self.logger = logging.getLogger(f"validation.{self.__class__.__name__}")

        # 生成されたファイルの追跡
        self.generated_files: list[str] = []
        self.generated_directories: list[str] = []

        # クイックモード設定
        self.quick_mode = False

        # サンプルテキストデータ
        self.sample_texts = {
            "ja": [
                "これは日本語のサンプルテキストです。",
                "DocMindアプリケーションのテスト用ドキュメントです。",
                "検索機能の動作確認を行います。",
                "全文検索とセマンティック検索の両方をテストします。",
                "パフォーマンスと精度の検証が重要です。",
            ],
            "en": [
                "This is a sample text in English.",
                "Test document for DocMind application.",
                "We are testing the search functionality.",
                "Both full-text and semantic search will be tested.",
                "Performance and accuracy validation is important.",
            ],
        }

        # サポートするファイル形式
        self.supported_formats = ["txt", "md", "json", "csv"]
        if HAS_DOCUMENT_LIBS:
            self.supported_formats.extend(["docx", "xlsx", "pdf"])

        self.logger.debug("テストデータ生成クラスを初期化しました")

    def setup_test_environment(self, base_dir: str) -> None:
        """
        テスト環境のセットアップ

        Args:
            base_dir: ベースディレクトリのパス
        """
        self.logger.info(f"テストデータ生成環境をセットアップします: {base_dir}")

        # ベースディレクトリの作成
        os.makedirs(base_dir, exist_ok=True)

        # クイックモードを有効にして高速化
        self.set_quick_mode(True)

        self.logger.debug("テストデータ生成環境のセットアップが完了しました")

    def set_quick_mode(self, enabled: bool) -> None:
        """
        クイックモードの設定

        Args:
            enabled: クイックモードを有効にするかどうか
        """
        self.quick_mode = enabled
        if enabled:
            self.logger.info("クイックモードが有効になりました（テストデータを削減）")
        else:
            self.logger.info("クイックモードが無効になりました")

    def generate_test_documents(self, count: int) -> list:
        """
        テスト用のDocumentオブジェクトを生成

        Args:
            count: 生成するドキュメント数

        Returns:
            Documentオブジェクトのリスト
        """
        # クイックモードの場合、ドキュメント数を削減
        if self.quick_mode:
            count = min(count, 5)  # さらに削減

        documents = []

        # DocMindのモデルをインポート
        import os
        import sys

        sys.path.append(os.path.join(os.path.dirname(__file__), "..", ".."))

        from src.data.models import Document

        for i in range(count):
            # 一時ファイルを作成
            temp_file = tempfile.NamedTemporaryFile(
                mode="w", suffix=".txt", delete=False, encoding="utf-8"
            )

            # テストコンテンツを生成
            content = self._generate_text_content(
                1.0 if self.quick_mode else 5.0,  # クイックモードでは小さなファイル
                "ja",
            )

            temp_file.write(content)
            temp_file.close()

            # Documentオブジェクトを作成
            doc = Document.create_from_file(temp_file.name, content)
            doc.id = f"test_doc_{i:04d}"
            doc.title = f"テストドキュメント {i}"

            documents.append(doc)
            self.generated_files.append(temp_file.name)

        return documents

    def generate_dataset(self, config: TestDatasetConfig) -> dict[str, Any]:
        """
        テストデータセットの生成

        Args:
            config: データセット生成設定

        Returns:
            生成結果の辞書
        """
        self.logger.info(
            f"テストデータセット '{config.dataset_name}' の生成を開始します"
        )

        # 出力ディレクトリの作成
        os.makedirs(config.output_directory, exist_ok=True)
        self.generated_directories.append(config.output_directory)

        # ファイル形式の決定
        file_types = config.file_types or self.supported_formats

        generated_files = []
        generation_stats = {
            "total_files": 0,
            "by_type": {},
            "total_size_mb": 0.0,
            "corrupted_files": 0,
            "large_files": 0,
            "special_char_files": 0,
        }

        # ファイル生成
        for i in range(config.file_count):
            file_type = random.choice(file_types)

            try:
                file_path, file_info = self._generate_single_file(
                    config.output_directory, file_type, config, i
                )

                generated_files.append(file_path)
                self.generated_files.append(file_path)

                # 統計更新
                generation_stats["total_files"] += 1
                generation_stats["by_type"][file_type] = (
                    generation_stats["by_type"].get(file_type, 0) + 1
                )
                generation_stats["total_size_mb"] += file_info["size_mb"]

                if file_info.get("is_corrupted"):
                    generation_stats["corrupted_files"] += 1
                if file_info.get("is_large"):
                    generation_stats["large_files"] += 1
                if file_info.get("has_special_chars"):
                    generation_stats["special_char_files"] += 1

            except Exception as e:
                self.logger.error(f"ファイル生成中にエラーが発生しました: {e}")

        # メタデータファイルの生成
        metadata_path = os.path.join(config.output_directory, "dataset_metadata.json")
        metadata = {
            "dataset_name": config.dataset_name,
            "generation_time": datetime.now().isoformat(),
            "config": {
                "file_count": config.file_count,
                "file_types": file_types,
                "size_range_kb": config.size_range_kb,
                "content_language": config.content_language,
                "include_corrupted": config.include_corrupted,
                "include_large_files": config.include_large_files,
                "include_special_chars": config.include_special_chars,
            },
            "statistics": generation_stats,
            "generated_files": generated_files,
        }

        with open(metadata_path, "w", encoding="utf-8") as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)

        self.generated_files.append(metadata_path)

        self.logger.info(
            f"テストデータセット生成完了: {generation_stats['total_files']}ファイル, "
            f"{generation_stats['total_size_mb']:.2f}MB"
        )

        return {
            "dataset_path": config.output_directory,
            "metadata_path": metadata_path,
            "statistics": generation_stats,
            "generated_files": generated_files,
        }

    def _generate_single_file(
        self, output_dir: str, file_type: str, config: TestDatasetConfig, index: int
    ) -> tuple[str, dict[str, Any]]:
        """
        単一ファイルの生成

        Args:
            output_dir: 出力ディレクトリ
            file_type: ファイル形式
            config: 生成設定
            index: ファイルインデックス

        Returns:
            (ファイルパス, ファイル情報)のタプル
        """
        # ファイル名の生成
        base_name = f"test_document_{index:04d}"

        # 特殊文字を含むファイル名の生成（設定に応じて）
        if config.include_special_chars and random.random() < 0.1:
            special_chars = "日本語ファイル名_テスト"
            base_name = f"{base_name}_{special_chars}"

        file_path = os.path.join(output_dir, f"{base_name}.{file_type}")

        # ファイルサイズの決定
        min_size, max_size = config.size_range_kb

        # クイックモードでは小さなファイルサイズに制限
        if self.quick_mode:
            max_size = min(max_size, 10)  # 最大10KB
            min_size = min(min_size, 1)  # 最小1KB

        if (
            config.include_large_files and random.random() < 0.05
        ):  # 5%の確率で大きなファイル
            if self.quick_mode:
                target_size_kb = random.randint(
                    max_size, max_size * 2
                )  # クイックモードでは控えめに
            else:
                target_size_kb = random.randint(max_size * 10, max_size * 100)
            is_large = True
        else:
            target_size_kb = random.randint(min_size, max_size)
            is_large = False

        # 破損ファイルの生成（設定に応じて）
        is_corrupted = config.include_corrupted and random.random() < 0.02  # 2%の確率

        # ファイル生成
        file_info = {
            "size_mb": 0.0,
            "is_corrupted": is_corrupted,
            "is_large": is_large,
            "has_special_chars": config.include_special_chars and "日本語" in base_name,
        }

        if file_type == "txt":
            self._generate_text_file(file_path, target_size_kb, config, is_corrupted)
        elif file_type == "md":
            self._generate_markdown_file(
                file_path, target_size_kb, config, is_corrupted
            )
        elif file_type == "json":
            self._generate_json_file(file_path, target_size_kb, config, is_corrupted)
        elif file_type == "csv":
            self._generate_csv_file(file_path, target_size_kb, config, is_corrupted)
        elif file_type == "docx" and HAS_DOCUMENT_LIBS:
            self._generate_docx_file(file_path, target_size_kb, config, is_corrupted)
        elif file_type == "xlsx" and HAS_DOCUMENT_LIBS:
            self._generate_xlsx_file(file_path, target_size_kb, config, is_corrupted)
        elif file_type == "pdf" and HAS_DOCUMENT_LIBS:
            self._generate_pdf_file(file_path, target_size_kb, config, is_corrupted)
        else:
            # フォールバック: テキストファイルとして生成
            self._generate_text_file(file_path, target_size_kb, config, is_corrupted)

        # ファイルサイズの取得
        if os.path.exists(file_path):
            file_info["size_mb"] = os.path.getsize(file_path) / (1024 * 1024)

        return file_path, file_info

    def _generate_text_file(
        self,
        file_path: str,
        target_size_kb: int,
        config: TestDatasetConfig,
        is_corrupted: bool,
    ) -> None:
        """テキストファイルの生成"""
        content = self._generate_text_content(target_size_kb, config.content_language)

        if is_corrupted:
            # 破損データの挿入
            content = (
                content[: len(content) // 2]
                + "\x00\xff\xfe"
                + content[len(content) // 2 :]
            )

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

    def _generate_markdown_file(
        self,
        file_path: str,
        target_size_kb: int,
        config: TestDatasetConfig,
        is_corrupted: bool,
    ) -> None:
        """Markdownファイルの生成"""
        content = "# テストドキュメント\n\n"
        content += "## 概要\n\n"
        content += self._generate_text_content(
            target_size_kb - 1, config.content_language
        )
        content += "\n\n## 詳細\n\n"
        content += "- リスト項目1\n- リスト項目2\n- リスト項目3\n"

        if is_corrupted:
            content = content.replace("##", "##\x00")

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

    def _generate_json_file(
        self,
        file_path: str,
        target_size_kb: int,
        config: TestDatasetConfig,
        is_corrupted: bool,
    ) -> None:
        """JSONファイルの生成"""
        data = {
            "title": "テストドキュメント",
            "content": self._generate_text_content(
                target_size_kb - 1, config.content_language
            ),
            "metadata": {
                "created": datetime.now().isoformat(),
                "language": config.content_language,
                "type": "test_document",
            },
        }

        if is_corrupted:
            # 不正なJSONを生成
            json_str = json.dumps(data, ensure_ascii=False, indent=2)
            json_str = json_str[:-10] + "invalid_json"
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(json_str)
        else:
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

    def _generate_csv_file(
        self,
        file_path: str,
        target_size_kb: int,
        config: TestDatasetConfig,
        is_corrupted: bool,
    ) -> None:
        """CSVファイルの生成"""
        rows = []
        header = ["ID", "タイトル", "内容", "作成日"]
        rows.append(header)

        # 目標サイズに達するまで行を追加
        row_count = max(5, min(50, target_size_kb // 2))  # 概算、最大50行に制限
        for i in range(row_count):
            row = [
                str(i),
                f"テストタイトル{i}",
                self._generate_text_content(0.1, config.content_language)[:100],
                (datetime.now() - timedelta(days=random.randint(0, 365))).strftime(
                    "%Y-%m-%d"
                ),
            ]
            rows.append(row)

        with open(file_path, "w", encoding="utf-8", newline="") as f:
            writer = csv.writer(f)
            if is_corrupted:
                # 不正なCSVを生成（カンマを削除）
                for row in rows[: len(rows) // 2]:
                    writer.writerow(row)
                f.write("invalid,csv,data\nwithout,proper,structure")
            else:
                writer.writerows(rows)

    def _generate_docx_file(
        self,
        file_path: str,
        target_size_kb: int,
        config: TestDatasetConfig,
        is_corrupted: bool,
    ) -> None:
        """Word文書ファイルの生成"""
        if not HAS_DOCUMENT_LIBS:
            return

        doc = Document()
        doc.add_heading("テストドキュメント", 0)

        content = self._generate_text_content(target_size_kb, config.content_language)
        paragraphs = content.split("\n")

        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph)

        if is_corrupted:
            # 破損ファイルの生成は困難なので、通常ファイルを生成後にバイナリ操作
            doc.save(file_path)
            with open(file_path, "r+b") as f:
                f.seek(100)
                f.write(b"\x00\xff\xfe")
        else:
            doc.save(file_path)

    def _generate_xlsx_file(
        self,
        file_path: str,
        target_size_kb: int,
        config: TestDatasetConfig,
        is_corrupted: bool,
    ) -> None:
        """Excelファイルの生成"""
        if not HAS_DOCUMENT_LIBS:
            return

        wb = Workbook()
        ws = wb.active
        ws.title = "テストデータ"

        # ヘッダー
        headers = ["ID", "タイトル", "内容", "数値", "日付"]
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)

        # データ行
        row_count = max(5, min(100, target_size_kb // 5))  # 最大100行に制限
        for row in range(2, row_count + 2):
            ws.cell(row=row, column=1, value=row - 1)
            ws.cell(row=row, column=2, value=f"テストタイトル{row-1}")
            ws.cell(
                row=row,
                column=3,
                value=self._generate_text_content(0.1, config.content_language)[:50],
            )
            ws.cell(row=row, column=4, value=random.randint(1, 1000))
            ws.cell(
                row=row,
                column=5,
                value=datetime.now() - timedelta(days=random.randint(0, 365)),
            )

        if is_corrupted:
            wb.save(file_path)
            # バイナリ操作で破損
            with open(file_path, "r+b") as f:
                f.seek(200)
                f.write(b"\x00\xff\xfe")
        else:
            wb.save(file_path)

    def _generate_pdf_file(
        self,
        file_path: str,
        target_size_kb: int,
        config: TestDatasetConfig,
        is_corrupted: bool,
    ) -> None:
        """PDFファイルの生成"""
        if not HAS_DOCUMENT_LIBS:
            return

        # PyMuPDFを使用してPDFを生成
        doc = fitz.open()
        page = doc.new_page()

        content = self._generate_text_content(target_size_kb, config.content_language)

        # テキストを挿入
        text_rect = fitz.Rect(50, 50, 550, 750)
        page.insert_textbox(text_rect, content, fontsize=12, fontname="helv")

        if is_corrupted:
            doc.save(file_path)
            doc.close()
            # バイナリ操作で破損
            with open(file_path, "r+b") as f:
                f.seek(50)
                f.write(b"\x00\xff\xfe")
        else:
            doc.save(file_path)
            doc.close()

    def _generate_text_content(self, target_size_kb: float, language: str) -> str:
        """指定サイズのテキストコンテンツを生成"""
        sample_texts = self.sample_texts.get(language, self.sample_texts["ja"])

        # 最小サイズを保証
        target_size_kb = max(0.1, target_size_kb)  # 最小100バイト
        target_size_bytes = int(target_size_kb * 1024)

        content = ""
        max_iterations = 1000  # 無限ループ防止
        iteration_count = 0

        while (
            len(content.encode("utf-8")) < target_size_bytes
            and iteration_count < max_iterations
        ):
            # ランダムにサンプルテキストを選択
            text = random.choice(sample_texts)

            # ランダムな単語を追加
            random_words = self._generate_random_words(language, random.randint(5, 20))
            text += " " + " ".join(random_words)

            content += text + "\n\n"
            iteration_count += 1

        # 文字単位で適切にトリミング
        content_bytes = content.encode("utf-8")
        if len(content_bytes) > target_size_bytes:
            # バイト単位で切り詰めて、文字境界で調整
            truncated = content_bytes[:target_size_bytes]
            try:
                content = truncated.decode("utf-8")
            except UnicodeDecodeError:
                # 文字境界で切れた場合は、安全な位置まで戻る
                for i in range(len(truncated) - 1, max(0, len(truncated) - 4), -1):
                    try:
                        content = truncated[:i].decode("utf-8")
                        break
                    except UnicodeDecodeError:
                        continue

        return content

    def _generate_random_words(self, language: str, count: int) -> list[str]:
        """ランダムな単語を生成"""
        if language == "ja":
            # 日本語の単語候補
            words = [
                "検索",
                "ドキュメント",
                "テスト",
                "機能",
                "システム",
                "データ",
                "ファイル",
                "処理",
                "結果",
                "情報",
                "管理",
                "設定",
                "画面",
                "操作",
                "確認",
            ]
        else:
            # 英語の単語候補
            words = [
                "search",
                "document",
                "test",
                "function",
                "system",
                "data",
                "file",
                "process",
                "result",
                "information",
                "management",
                "setting",
                "screen",
                "operation",
                "check",
            ]

        return [random.choice(words) for _ in range(count)]

    def generate_large_dataset(
        self, output_dir: str, document_count: int = 50000
    ) -> dict[str, Any]:
        """
        大規模テストデータセットの生成

        Args:
            output_dir: 出力ディレクトリ
            document_count: 生成するドキュメント数

        Returns:
            生成結果の辞書
        """
        config = TestDatasetConfig(
            dataset_name="large_dataset",
            output_directory=output_dir,
            file_count=document_count,
            size_range_kb=(1, 100),
            include_large_files=True,
            include_special_chars=True,
        )

        return self.generate_dataset(config)

    def generate_edge_case_dataset(self, output_dir: str) -> dict[str, Any]:
        """
        エッジケーステストデータセットの生成

        Args:
            output_dir: 出力ディレクトリ

        Returns:
            生成結果の辞書
        """
        config = TestDatasetConfig(
            dataset_name="edge_case_dataset",
            output_directory=output_dir,
            file_count=100,
            size_range_kb=(0, 10000),  # 0KBから10MBまで
            include_corrupted=True,
            include_large_files=True,
            include_special_chars=True,
        )

        return self.generate_dataset(config)

    def cleanup(self) -> None:
        """生成されたテストデータのクリーンアップ"""
        self.logger.info("テストデータのクリーンアップを開始します")

        # 生成されたファイルの削除
        for file_path in self.generated_files:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    self.logger.debug(f"ファイルを削除しました: {file_path}")
            except Exception as e:
                self.logger.warning(f"ファイル削除に失敗しました: {file_path} - {e}")

        # 生成されたディレクトリの削除
        for dir_path in self.generated_directories:
            try:
                if os.path.exists(dir_path) and os.path.isdir(dir_path):
                    shutil.rmtree(dir_path)
                    self.logger.debug(f"ディレクトリを削除しました: {dir_path}")
            except Exception as e:
                self.logger.warning(f"ディレクトリ削除に失敗しました: {dir_path} - {e}")

        self.generated_files.clear()
        self.generated_directories.clear()

        self.logger.info("テストデータのクリーンアップが完了しました")
