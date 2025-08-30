"""
DocMindアプリケーション用のドキュメント処理システム

このモジュールは、様々なファイル形式からテキストを抽出するDocumentProcessorクラスを提供します。
サポートされる形式：PDF、Word、Excel、Markdown、テキストファイル
"""

import logging
import os
from pathlib import Path
from typing import Any

import chardet

# ドキュメント処理ライブラリ
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

# 内部モジュール
from ..data.models import Document, FileType
from ..utils.exceptions import DocumentProcessingError


class DocumentProcessor:
    """様々なファイル形式からテキストを抽出するプロセッサークラス

    このクラスは、PDF、Word、Excel、Markdown、テキストファイルから
    テキストコンテンツを抽出し、Documentオブジェクトを作成します。
    各ファイル形式に対して適切なライブラリを使用し、エラーハンドリングを提供します。
    """

    def __init__(self):
        """DocumentProcessorを初期化"""
        self.logger = logging.getLogger(__name__)
        self._check_dependencies()

    def _check_dependencies(self):
        """必要な依存関係がインストールされているかチェック"""
        missing_deps = []

        if fitz is None:
            missing_deps.append("PyMuPDF (PDF処理用)")

        if DocxDocument is None:
            missing_deps.append("python-docx (Word文書処理用)")

        if load_workbook is None:
            missing_deps.append("openpyxl (Excel処理用)")

        if missing_deps:
            self.logger.warning(f"以下の依存関係が不足しています: {', '.join(missing_deps)}")

    def process_file(self, file_path: str) -> Document:
        """ファイルを処理してDocumentオブジェクトを作成

        Args:
            file_path (str): 処理するファイルのパス

        Returns:
            Document: 作成されたDocumentオブジェクト

        Raises:
            DocumentProcessingError: ファイル処理中にエラーが発生した場合
        """
        if not os.path.exists(file_path):
            raise DocumentProcessingError(
                f"ファイルが存在しません: {file_path}",
                file_path=file_path
            )

        try:
            # ファイルタイプを判定
            file_type = FileType.from_extension(file_path)
            self.logger.info(f"ファイル処理開始: {file_path} (タイプ: {file_type.value})")

            # ファイルタイプに応じてテキストを抽出
            content = self._extract_text_by_type(file_path, file_type)

            # Documentオブジェクトを作成
            document = Document.create_from_file(file_path, content)

            self.logger.info(f"ファイル処理完了: {file_path} (コンテンツ長: {len(content)}文字)")
            return document

        except Exception as e:
            error_msg = f"ファイル処理中にエラーが発生しました: {file_path}"
            self.logger.error(f"{error_msg} - {str(e)}")
            raise DocumentProcessingError(
                error_msg,
                file_path=file_path,
                file_type=file_type.value if 'file_type' in locals() else "unknown",
                details=str(e)
            )

    def _extract_text_by_type(self, file_path: str, file_type: FileType) -> str:
        """ファイルタイプに応じてテキストを抽出

        Args:
            file_path (str): ファイルパス
            file_type (FileType): ファイルタイプ

        Returns:
            str: 抽出されたテキスト
        """
        if file_type == FileType.PDF:
            return self.extract_pdf_text(file_path)
        elif file_type == FileType.WORD:
            return self.extract_word_text(file_path)
        elif file_type == FileType.EXCEL:
            return self.extract_excel_text(file_path)
        elif file_type == FileType.MARKDOWN:
            return self.extract_markdown_text(file_path)
        elif file_type == FileType.TEXT:
            return self.extract_text_file(file_path)
        else:
            # 未知のファイルタイプの場合、テキストファイルとして処理を試行
            self.logger.warning(f"未知のファイルタイプ: {file_type.value}, テキストファイルとして処理します")
            return self.extract_text_file(file_path)

    def extract_pdf_text(self, file_path: str) -> str:
        """PDFファイルからテキストを抽出

        PyMuPDFを使用してPDFファイルからテキストコンテンツを抽出します。

        Args:
            file_path (str): PDFファイルのパス

        Returns:
            str: 抽出されたテキスト

        Raises:
            DocumentProcessingError: PDF処理中にエラーが発生した場合
        """
        if fitz is None:
            raise DocumentProcessingError(
                "PyMuPDFがインストールされていません。PDF処理には PyMuPDF が必要です。",
                file_path=file_path,
                file_type="pdf"
            )

        try:
            text_content = []

            # PDFドキュメントを開く
            with fitz.open(file_path) as pdf_doc:
                self.logger.debug(f"PDF処理開始: {file_path} (ページ数: {len(pdf_doc)})")

                # 各ページからテキストを抽出
                for page_num in range(len(pdf_doc)):
                    try:
                        page = pdf_doc[page_num]
                        page_text = page.get_text()

                        if page_text.strip():
                            text_content.append(page_text)
                            self.logger.debug(f"ページ {page_num + 1}: {len(page_text)}文字抽出")

                    except Exception as e:
                        self.logger.warning(f"ページ {page_num + 1} の処理中にエラー: {str(e)}")
                        continue

            extracted_text = "\n\n".join(text_content)

            if not extracted_text.strip():
                self.logger.warning(f"PDFからテキストが抽出されませんでした: {file_path}")
                return ""

            self.logger.debug(f"PDF処理完了: {file_path} (総文字数: {len(extracted_text)})")
            return extracted_text

        except Exception as e:
            raise DocumentProcessingError(
                f"PDF処理中にエラーが発生しました: {str(e)}",
                file_path=file_path,
                file_type="pdf",
                details=str(e)
            )

    def extract_word_text(self, file_path: str) -> str:
        """Wordドキュメントからテキストを抽出

        python-docxを使用してWordファイルからテキストコンテンツを抽出します。

        Args:
            file_path (str): Wordファイルのパス

        Returns:
            str: 抽出されたテキスト

        Raises:
            DocumentProcessingError: Word処理中にエラーが発生した場合
        """
        if DocxDocument is None:
            raise DocumentProcessingError(
                "python-docxがインストールされていません。Word文書処理には python-docx が必要です。",
                file_path=file_path,
                file_type="word"
            )

        try:
            self.logger.debug(f"Word文書処理開始: {file_path}")

            # Wordドキュメントを開く
            doc = DocxDocument(file_path)
            text_content = []

            # 段落からテキストを抽出
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # テーブルからテキストを抽出
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            extracted_text = "\n".join(text_content)

            if not extracted_text.strip():
                self.logger.warning(f"Word文書からテキストが抽出されませんでした: {file_path}")
                return ""

            self.logger.debug(f"Word文書処理完了: {file_path} (総文字数: {len(extracted_text)})")
            return extracted_text

        except Exception as e:
            raise DocumentProcessingError(
                f"Word文書処理中にエラーが発生しました: {str(e)}",
                file_path=file_path,
                file_type="word",
                details=str(e)
            )

    def extract_excel_text(self, file_path: str) -> str:
        """Excelファイルからテキストを抽出

        openpyxlを使用してExcelファイルのすべてのシートとセルからテキストを抽出します。

        Args:
            file_path (str): Excelファイルのパス

        Returns:
            str: 抽出されたテキスト

        Raises:
            DocumentProcessingError: Excel処理中にエラーが発生した場合
        """
        if load_workbook is None:
            raise DocumentProcessingError(
                "openpyxlがインストールされていません。Excel処理には openpyxl が必要です。",
                file_path=file_path,
                file_type="excel"
            )

        try:
            self.logger.debug(f"Excel処理開始: {file_path}")

            # Excelワークブックを開く
            workbook = load_workbook(file_path, read_only=True, data_only=True)
            text_content = []

            # 各シートを処理
            for sheet_name in workbook.sheetnames:
                try:
                    sheet = workbook[sheet_name]
                    sheet_content = []

                    self.logger.debug(f"シート処理開始: {sheet_name}")

                    # シート名を追加
                    sheet_content.append(f"=== シート: {sheet_name} ===")

                    # 各行を処理
                    for row in sheet.iter_rows(values_only=True):
                        row_text = []
                        for cell_value in row:
                            if cell_value is not None:
                                # セル値を文字列に変換
                                cell_str = str(cell_value).strip()
                                if cell_str:
                                    row_text.append(cell_str)

                        if row_text:
                            sheet_content.append(" | ".join(row_text))

                    if len(sheet_content) > 1:  # シート名以外にコンテンツがある場合
                        text_content.extend(sheet_content)
                        self.logger.debug(f"シート {sheet_name}: {len(sheet_content) - 1}行処理")

                except Exception as e:
                    self.logger.warning(f"シート {sheet_name} の処理中にエラー: {str(e)}")
                    continue

            workbook.close()

            extracted_text = "\n".join(text_content)

            if not extracted_text.strip():
                self.logger.warning(f"Excelファイルからテキストが抽出されませんでした: {file_path}")
                return ""

            self.logger.debug(f"Excel処理完了: {file_path} (総文字数: {len(extracted_text)})")
            return extracted_text

        except Exception as e:
            raise DocumentProcessingError(
                f"Excel処理中にエラーが発生しました: {str(e)}",
                file_path=file_path,
                file_type="excel",
                details=str(e)
            )

    def extract_markdown_text(self, file_path: str) -> str:
        """Markdownファイルからテキストを抽出

        Markdownファイルから構造を保持しながらテキストコンテンツを抽出します。

        Args:
            file_path (str): Markdownファイルのパス

        Returns:
            str: 抽出されたテキスト

        Raises:
            DocumentProcessingError: Markdown処理中にエラーが発生した場合
        """
        try:
            self.logger.debug(f"Markdown処理開始: {file_path}")

            # エンコーディングを検出してファイルを読み込み
            encoding = self._detect_encoding(file_path)

            with open(file_path, encoding=encoding) as file:
                content = file.read()

            if not content.strip():
                self.logger.warning(f"Markdownファイルが空です: {file_path}")
                return ""

            # 基本的なMarkdown記法を処理（見出し、リスト、コードブロックなど）
            processed_content = self._process_markdown_content(content)

            self.logger.debug(f"Markdown処理完了: {file_path} (総文字数: {len(processed_content)})")
            return processed_content

        except Exception as e:
            raise DocumentProcessingError(
                f"Markdown処理中にエラーが発生しました: {str(e)}",
                file_path=file_path,
                file_type="markdown",
                details=str(e)
            )

    def extract_text_file(self, file_path: str) -> str:
        """テキストファイルからテキストを抽出

        エンコーディング検出を使用してテキストファイルを読み込みます。

        Args:
            file_path (str): テキストファイルのパス

        Returns:
            str: 抽出されたテキスト

        Raises:
            DocumentProcessingError: テキストファイル処理中にエラーが発生した場合
        """
        try:
            self.logger.debug(f"テキストファイル処理開始: {file_path}")

            # エンコーディングを検出してファイルを読み込み
            encoding = self._detect_encoding(file_path)

            with open(file_path, encoding=encoding) as file:
                content = file.read()

            if not content.strip():
                self.logger.warning(f"テキストファイルが空です: {file_path}")
                return ""

            self.logger.debug(f"テキストファイル処理完了: {file_path} (総文字数: {len(content)})")
            return content

        except Exception as e:
            raise DocumentProcessingError(
                f"テキストファイル処理中にエラーが発生しました: {str(e)}",
                file_path=file_path,
                file_type="text",
                details=str(e)
            )

    def _detect_encoding(self, file_path: str) -> str:
        """ファイルのエンコーディングを検出

        Args:
            file_path (str): ファイルパス

        Returns:
            str: 検出されたエンコーディング
        """
        try:
            # ファイルの一部を読み込んでエンコーディングを検出
            with open(file_path, 'rb') as file:
                raw_data = file.read(10000)  # 最初の10KBを読み込み

            if raw_data:
                detected = chardet.detect(raw_data)
                encoding = detected.get('encoding', 'utf-8')
                confidence = detected.get('confidence', 0)

                self.logger.debug(f"エンコーディング検出: {encoding} (信頼度: {confidence:.2f})")

                # 信頼度が低い場合はUTF-8を使用
                if confidence < 0.7:
                    self.logger.debug("信頼度が低いため UTF-8 を使用します")
                    encoding = 'utf-8'

                return encoding
            else:
                return 'utf-8'

        except Exception as e:
            self.logger.warning(f"エンコーディング検出に失敗: {str(e)}, UTF-8を使用します")
            return 'utf-8'

    def _process_markdown_content(self, content: str) -> str:
        """Markdownコンテンツを処理

        基本的なMarkdown記法を処理して、検索しやすい形式に変換します。

        Args:
            content (str): 元のMarkdownコンテンツ

        Returns:
            str: 処理されたコンテンツ
        """
        lines = content.split('\n')
        processed_lines = []

        for line in lines:
            line = line.strip()

            if not line:
                continue

            # 見出し記法を処理
            if line.startswith('#'):
                # # を削除して見出しテキストを抽出
                heading_text = line.lstrip('#').strip()
                if heading_text:
                    processed_lines.append(f"見出し: {heading_text}")

            # リスト記法を処理
            elif line.startswith(('- ', '* ', '+ ')) or line.lstrip().startswith(('- ', '* ', '+ ')):
                list_text = line.lstrip('- *+').strip()
                if list_text:
                    processed_lines.append(f"リスト項目: {list_text}")

            # 番号付きリストを処理
            elif line.lstrip() and line.lstrip()[0].isdigit() and '. ' in line:
                parts = line.split('. ', 1)
                if len(parts) > 1:
                    list_text = parts[1].strip()
                    if list_text:
                        processed_lines.append(f"番号付きリスト: {list_text}")

            # コードブロックは除外（```で囲まれた部分）
            elif line.startswith('```'):
                continue

            # インラインコードやリンクなどの記法を除去
            else:
                # 基本的なMarkdown記法を除去
                cleaned_line = line

                # **太字** を除去
                import re
                cleaned_line = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned_line)

                # *斜体* を除去
                cleaned_line = re.sub(r'\*(.*?)\*', r'\1', cleaned_line)

                # `インラインコード` を除去
                cleaned_line = re.sub(r'`(.*?)`', r'\1', cleaned_line)

                # [リンクテキスト](URL) からリンクテキストのみ抽出
                cleaned_line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned_line)

                if cleaned_line.strip():
                    processed_lines.append(cleaned_line.strip())

        return '\n'.join(processed_lines)

    def get_supported_extensions(self) -> dict[str, FileType]:
        """サポートされているファイル拡張子とFileTypeのマッピングを取得

        Returns:
            Dict[str, FileType]: 拡張子とFileTypeのマッピング
        """
        return {
            '.pdf': FileType.PDF,
            '.doc': FileType.WORD,
            '.docx': FileType.WORD,
            '.xls': FileType.EXCEL,
            '.xlsx': FileType.EXCEL,
            '.md': FileType.MARKDOWN,
            '.markdown': FileType.MARKDOWN,
            '.txt': FileType.TEXT,
            '.text': FileType.TEXT
        }

    def is_supported_file(self, file_path: str) -> bool:
        """ファイルがサポートされているかチェック

        Args:
            file_path (str): チェックするファイルのパス

        Returns:
            bool: サポートされている場合True
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.get_supported_extensions()

    def get_file_info(self, file_path: str) -> dict[str, Any]:
        """ファイルの基本情報を取得

        Args:
            file_path (str): ファイルパス

        Returns:
            Dict[str, Any]: ファイル情報
        """
        if not os.path.exists(file_path):
            return {}

        try:
            file_stat = os.stat(file_path)
            path_obj = Path(file_path)

            return {
                'name': path_obj.name,
                'stem': path_obj.stem,
                'suffix': path_obj.suffix,
                'size': file_stat.st_size,
                'created': file_stat.st_ctime,
                'modified': file_stat.st_mtime,
                'file_type': FileType.from_extension(file_path),
                'is_supported': self.is_supported_file(file_path)
            }
        except Exception as e:
            self.logger.error(f"ファイル情報取得エラー: {file_path} - {str(e)}")
            return {}
